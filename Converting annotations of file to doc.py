#!/usr/bin/env python
# coding: utf-8

# In[1]:


get_ipython().system('pip install pymupdf python-docx')


# In[7]:





# In[2]:


import fitz  # PyMuPDF
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

# ---------------- PDF Processing Function ----------------
def export_annotations(pdf_path, output_path, include_author, include_type):
    doc = fitz.open(pdf_path)
    word = Document()
    word.add_heading("PDF Annotations (Ordered)", level=1)

    counter = 1

    for page_index in range(len(doc)):
        page = doc[page_index]
        annot = page.first_annot

        while annot:
            info = annot.info
            annot_type = annot.type[1]
            author = info.get("title", "Unknown")
            comment = info.get("content", "").strip()
            highlighted_text = ""

            # ---- Extract highlighted text ----
            if annot_type == "Highlight":
                quad_points = annot.vertices
                if quad_points:
                    text_parts = []
                    # Iterate through the quad points (usually 4 per highlight span)
                    for i in range(0, len(quad_points), 4):
                        quad = quad_points[i:i+4]
                        rect = fitz.Quad(quad).rect
                        text_parts.append(page.get_text("text", clip=rect))
                    
                    highlighted_text = " ".join(text_parts).strip()

            # Determine the main content for the annotation
            main_content = ""
            if highlighted_text:
                main_content = highlighted_text
            elif comment:
                # Capture the comment for annotations like 'Text' or 'Note'
                main_content = comment
            
            # Only proceed if there is actual content
            if not main_content:
                annot = annot.next
                continue

            # ---- Write to Word with the new format ----
            
            # New combined output line: "Type: Page Number: Annotation Content"
            output_line = f"{annot_type if include_type else ''}{':' if include_type else ''} Page {page_index + 1}: {main_content}"
            word.add_paragraph(f"{counter}. {output_line}")
            
            # --- Add optional fields as a sub-entry ---
            # If the annotation type was already added to the main line, skip it here
            
            if include_author:
                word.add_paragraph(f"Author: {author}")

            # Add a separator and increment counter
            word.add_paragraph("-" * 40)
            counter += 1

            annot = annot.next

    word.save(output_path)

# ---------------- GUI Functions ----------------
def select_pdf():
    file_path = filedialog.askopenfilename(
        title="Select PDF",
        filetypes=[("PDF Files", "*.pdf")]
    )
    if file_path:
        pdf_path_var.set(file_path)

def export():
    pdf_path = pdf_path_var.get()
    if not pdf_path:
        messagebox.showwarning("No PDF", "Please select a PDF file.")
        return

    output_path = filedialog.asksaveasfilename(
        title="Save Word File",
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )

    if not output_path:
        return

    try:
        export_annotations(
            pdf_path,
            output_path,
            include_author_var.get(),
            include_type_var.get()
        )
        messagebox.showinfo("Success", "Annotations exported successfully!")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# ---------------- GUI Layout ----------------
root = tk.Tk()
root.title("PDF Annotation Exporter")
root.geometry("600x260")
root.resizable(False, False)

pdf_path_var = tk.StringVar()
include_author_var = tk.BooleanVar(value=True)
include_type_var = tk.BooleanVar(value=True)

tk.Label(root, text="Selected PDF:", font=("Arial", 10)).pack(pady=5)
tk.Entry(root, textvariable=pdf_path_var, width=70).pack(pady=5)

tk.Button(root, text="Browse PDF", command=select_pdf).pack(pady=5)

# ---- Options Frame ----
options_frame = tk.LabelFrame(root, text="Export Options", padx=10, pady=5)
options_frame.pack(pady=10)

tk.Checkbutton(
    options_frame,
    text="Require Author name",
    variable=include_author_var
).pack(anchor="w")

tk.Checkbutton(
    options_frame,
    text="Require Annotation Type",
    variable=include_type_var
).pack(anchor="w")

tk.Button(root, text="Export Annotations to Word", command=export).pack(pady=10)

root.mainloop()


# In[ ]:




